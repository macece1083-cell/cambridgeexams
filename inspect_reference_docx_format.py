from pathlib import Path
from collections import Counter
from docx import Document


FILES = [
    Path(r"C:\Users\User\Desktop\flyers 1 exam.docx"),
    Path(r"C:\Users\User\Desktop\impact 1 ket exam 1.docx"),
    Path(r"C:\Users\User\Desktop\Impact 2 KET Exam 1 - easier part 2.docx"),
    Path(r"C:\Users\User\Desktop\Impact 2 KET Exam 1.docx"),
]


def pt(value):
    if value is None:
        return None
    try:
        return round(value.pt, 2)
    except Exception:
        return None


def inch(value):
    if value is None:
        return None
    try:
        return round(value.inches, 2)
    except Exception:
        return None


def paragraph_signature(paragraph):
    style = paragraph.style.name if paragraph.style else ""
    text = paragraph.text.strip()
    runs = [r for r in paragraph.runs if r.text.strip()]
    fonts = Counter()
    sizes = Counter()
    bolds = Counter()
    for run in runs:
        font_name = run.font.name or ""
        font_size = pt(run.font.size)
        fonts[font_name] += 1
        sizes[font_size] += 1
        bolds[bool(run.bold)] += 1
    pf = paragraph.paragraph_format
    return {
        "style": style,
        "text": text[:140],
        "font": fonts.most_common(1)[0][0] if fonts else "",
        "size": sizes.most_common(1)[0][0] if sizes else None,
        "bold": bolds.most_common(1)[0][0] if bolds else False,
        "space_before": pt(pf.space_before),
        "space_after": pt(pf.space_after),
        "line_spacing": pf.line_spacing,
        "alignment": str(paragraph.alignment),
    }


def describe(path):
    doc = Document(path)
    sec = doc.sections[0]
    print(f"\n===== {path.name} =====")
    print(
        "page",
        inch(sec.page_width),
        "x",
        inch(sec.page_height),
        "margins",
        {
            "top": inch(sec.top_margin),
            "bottom": inch(sec.bottom_margin),
            "left": inch(sec.left_margin),
            "right": inch(sec.right_margin),
        },
    )
    style_counts = Counter(p.style.name if p.style else "" for p in doc.paragraphs if p.text.strip())
    print("styles", style_counts.most_common(10))
    font_counts = Counter()
    size_counts = Counter()
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text.strip():
                font_counts[r.font.name or ""] += 1
                size_counts[pt(r.font.size)] += 1
    print("fonts", font_counts.most_common(8))
    print("sizes", size_counts.most_common(8))
    print("tables", len(doc.tables))
    for i, table in enumerate(doc.tables[:5], 1):
        rows = len(table.rows)
        cols = len(table.columns)
        sample = " | ".join(cell.text.strip().replace("\n", " ")[:30] for cell in table.rows[0].cells[:4]) if rows else ""
        print(f"table {i}: {rows}x{cols} sample={sample!r}")
    print("first paragraphs")
    shown = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        sig = paragraph_signature(p)
        print(sig)
        shown += 1
        if shown >= 22:
            break


for file in FILES:
    if file.exists():
        describe(file)
    else:
        print(f"missing: {file}")
